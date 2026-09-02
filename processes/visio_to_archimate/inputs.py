"""Workload INPUTS — a system diagram plus optional requirements documents — by reference or path.

Where the BA reads from: historically a local filesystem path (the container's disk), which a
cloud workload cannot be handed. Inputs are therefore addressed like every other lab artifact:
an `art://<id>/<name>` reference in the shared artifact store (`shared/artifacts.py`, Postgres),
uploaded once with `python -m processes.visio_to_archimate.inputs upload <files...>`. Plain
paths still work for local runs. Parsing happens HERE, locally, with no egress: a `.vsdx` via the
visio-reader script, a `.docx`/`.pdf`/text document into plain text; an IMAGE is not parsed at
all — it is attached inline to the BA's message (kimi-k3 has vision through the gateway, verified).
Whatever text these produce then egresses to the model through the gateway (PII-guarded).
"""
import io
import os
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))
from shared import artifacts  # noqa: E402

IMAGE_TYPES = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
               ".gif": "image/gif", ".webp": "image/webp"}
DOC_TYPES = {".docx", ".pdf", ".md", ".markdown", ".txt", ".rst", ".csv"}
MAX_DOC_CHARS = int(os.environ.get("BA_MAX_DOC_CHARS", "60000"))   # keeps one doc inside a turn


def is_ref(src: str) -> bool:
    return isinstance(src, str) and src.startswith("art://")


def name_of(src: str) -> str:
    return src.rstrip("/").split("/")[-1]


def kind(src: str) -> str:
    """'vsdx' | 'image' | 'document' | 'unknown' — decided by the file name, ref or path alike."""
    ext = Path(name_of(src)).suffix.lower()
    if ext == ".vsdx":
        return "vsdx"
    if ext in IMAGE_TYPES:
        return "image"
    if ext in DOC_TYPES:
        return "document"
    return "unknown"


def media_type(src: str) -> str:
    return IMAGE_TYPES.get(Path(name_of(src)).suffix.lower(), "application/octet-stream")


def load(src: str) -> bytes:
    return artifacts.store().get(src) if is_ref(src) else Path(src).read_bytes()


def local_path(src: str) -> str:
    """A filesystem path for parsers that insist on one (the vsdx reader): a ref is materialised
    to a temp file; a path is returned as-is."""
    if not is_ref(src):
        return src
    p = Path(tempfile.gettempdir()) / name_of(src)
    p.write_bytes(load(src))
    return str(p)


def read_document(src: str) -> str:
    """Requirements document -> plain text (.docx paragraphs + tables, .pdf page text, else UTF-8)."""
    data, ext = load(src), Path(name_of(src)).suffix.lower()
    if ext == ".docx":
        import docx
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        text = "\n".join(parts)
    elif ext == ".pdf":
        from pypdf import PdfReader
        text = "\n\n".join((pg.extract_text() or "") for pg in PdfReader(io.BytesIO(data)).pages)
    else:
        text = data.decode("utf-8", errors="replace")
    if len(text) > MAX_DOC_CHARS:
        text = text[:MAX_DOC_CHARS] + f"\n\n[truncated: {len(text)} chars in total, first {MAX_DOC_CHARS} shown]"
    return text


MAX_EMBEDDED_IMAGES = int(os.environ.get("BA_MAX_EMBEDDED_IMAGES", "8"))
MIN_IMAGE_BYTES = 2048            # skip bullets/icons
MAX_IMAGE_EDGE = 1600             # downscale huge figures so one message stays cheap


def _normalise_image(data: bytes, ctype: str) -> tuple[bytes, str] | None:
    """Drop tiny decorations; convert/downscale to PNG/JPEG the vision model accepts."""
    if len(data) < MIN_IMAGE_BYTES:
        return None
    try:
        from PIL import Image
        im = Image.open(io.BytesIO(data))
        if min(im.size) < 64:
            return None
        if max(im.size) > MAX_IMAGE_EDGE or ctype not in ("image/png", "image/jpeg"):
            im.thumbnail((MAX_IMAGE_EDGE, MAX_IMAGE_EDGE))
            buf = io.BytesIO()
            im.convert("RGB").save(buf, "PNG")
            return buf.getvalue(), "image/png"
        return data, ctype
    except Exception:
        return None                  # not an image PIL understands (emf/wmf …) — skip


def extract_images(src: str) -> list[tuple[str, bytes, str]]:
    """Figures EMBEDDED in a requirements document -> [(label, bytes, media_type)], so the BA can
    read them with vision: a .docx's image parts, a .pdf's page images. Deterministic, local."""
    data, ext, out = load(src), Path(name_of(src)).suffix.lower(), []
    doc_name = name_of(src)
    if ext == ".docx":
        import docx
        d = docx.Document(io.BytesIO(data))
        for rel in d.part.rels.values():
            if "image" in rel.reltype and getattr(rel, "target_part", None) is not None:
                norm = _normalise_image(rel.target_part.blob, rel.target_part.content_type)
                if norm:
                    out.append((f"figure {len(out) + 1} embedded in {doc_name}", *norm))
            if len(out) >= MAX_EMBEDDED_IMAGES:
                break
    elif ext == ".pdf":
        from pypdf import PdfReader
        for pno, page in enumerate(PdfReader(io.BytesIO(data)).pages, 1):
            for img in page.images:
                norm = _normalise_image(img.data, "image/" + (img.name.rsplit(".", 1)[-1].lower() or "png"))
                if norm:
                    out.append((f"figure {len(out) + 1} embedded in {doc_name} (page {pno})", *norm))
                if len(out) >= MAX_EMBEDDED_IMAGES:
                    return out
    return out


def upload(paths) -> list[str]:
    """Store local files as artifacts; returns their art:// refs (what a cloud workload is given)."""
    return [artifacts.put_file(p, content_type=media_type(p) if kind(p) == "image" else None)
            for p in paths]


if __name__ == "__main__":
    if len(sys.argv) > 2 and sys.argv[1] == "upload":
        for p, ref in zip(sys.argv[2:], upload(sys.argv[2:])):
            print(f"{ref}\t{p}")
    else:
        sys.exit("usage: python -m processes.visio_to_archimate.inputs upload <diagram|doc> ...")
