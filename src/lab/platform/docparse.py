"""Deterministic parsing of a workload's INPUTS — bytes in, plain data out. No I/O, no egress.

Shared by two callers with different trust: `src/lab/substrate/mcp/storage/server.py` (the governed path — an
agent asks the gateway, the gateway asks storage-mcp, storage-mcp reads the bucket and parses here)
and `src/lab/workloads/visio_to_archimate/inputs.py` (local-dev paths only). Keeping the parsing here means
both produce byte-identical results and the sizing contract the visio-reader skill documents is
enforced in exactly one place.
"""
from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path

from lab.platform import filetypes
from lab.platform.contracts import split_fragment  # noqa: F401  (re-exported: the ONE `#page` parser)

# Views of the ONE file-type table (lab.platform.filetypes.FILE_TYPES): which extensions are read with
# vision and which as requirements text. Extend the table there, not these.
IMAGE_TYPES = {"." + ext: ct for ext, (ct, k) in filetypes.FILE_TYPES.items() if k == "image"}
DOC_TYPES = {"." + ext for ext, (_ct, k) in filetypes.FILE_TYPES.items() if k == "document"}
MAX_DOC_CHARS = int(os.environ.get("BA_MAX_DOC_CHARS", "60000"))      # keeps one doc inside a turn
MAX_EMBEDDED_IMAGES = int(os.environ.get("BA_MAX_EMBEDDED_IMAGES", "8"))
MIN_IMAGE_BYTES = 2048            # below this it is a bullet/icon, not a figure
MIN_IMAGE_EDGE = 64
MAX_IMAGE_EDGE = 1600             # the sizing contract: no image reaches a model larger than this


def _base(name: str) -> str:
    """File name with any `#page` fragment removed, for extension detection."""
    return split_fragment(name)[0].rstrip("/").split("/")[-1]


def kind(name: str) -> str:
    """'vsdx' | 'image' | 'document' | 'unknown' — decided by the file name (ref or path alike;
    a `#page` fragment is ignored)."""
    ext = Path(_base(name)).suffix.lower()
    if ext == ".vsdx":
        return "vsdx"
    if ext in IMAGE_TYPES:
        return "image"
    if ext in DOC_TYPES:
        return "document"
    return "unknown"


def media_type(name: str) -> str:
    """The media type an IMAGE input is attached with (octet-stream for anything that is not one
    of IMAGE_TYPES — documents and .vsdx are never sent as inline media)."""
    return IMAGE_TYPES.get(Path(_base(name)).suffix.lower(), "application/octet-stream")


def ext_of(name: str) -> str:
    return Path(_base(name)).suffix.lower()


def document_text(data: bytes, ext: str, max_chars: int = MAX_DOC_CHARS) -> str:
    """Requirements document -> plain text (.docx paragraphs + tables, .pdf page text, else UTF-8)."""
    if ext == ".docx":
        import docx
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        text = "\n".join(parts)
    elif ext == ".pdf":
        from pypdf import PdfReader
        text = "\n\n".join((pg.extract_text() or "") for pg in PdfReader(io.BytesIO(data)).pages)
    else:
        text = data.decode("utf-8", errors="replace")
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[truncated: {len(text)} chars in total, first {max_chars} shown]"
    return text


def normalise_image(data: bytes, ctype: str, max_edge: int = MAX_IMAGE_EDGE):
    """Apply the sizing contract: reject decorations (< MIN_IMAGE_BYTES or < MIN_IMAGE_EDGE px),
    downscale to max_edge, emit PNG unless already PNG/JPEG within bounds.
    Returns (bytes, media_type, (w, h)) or None when the input is not a usable figure."""
    if len(data) < MIN_IMAGE_BYTES:
        return None
    try:
        from PIL import Image
        im = Image.open(io.BytesIO(data))
        if min(im.size) < MIN_IMAGE_EDGE:
            return None
        if max(im.size) > max_edge or ctype not in ("image/png", "image/jpeg"):
            im.thumbnail((max_edge, max_edge))
            buf = io.BytesIO()
            im.convert("RGB").save(buf, "PNG")
            return buf.getvalue(), "image/png", im.size
        return data, ctype, im.size
    except Exception:
        return None                  # emf/wmf and friends — PIL cannot read them; skip


def embedded_figures(data: bytes, ext: str, doc_name: str, max_images: int = MAX_EMBEDDED_IMAGES,
                     max_edge: int = MAX_IMAGE_EDGE) -> list[tuple[str, bytes, str]]:
    """Figures EMBEDDED in a document -> [(label, bytes, media_type)]: a .docx's image parts, a
    .pdf's page images — normalised so a vision model can read them as diagrams/screenshots."""
    out: list[tuple[str, bytes, str]] = []
    if ext == ".docx":
        import docx
        d = docx.Document(io.BytesIO(data))
        for rel in d.part.rels.values():
            if "image" in rel.reltype and getattr(rel, "target_part", None) is not None:
                norm = normalise_image(rel.target_part.blob, rel.target_part.content_type, max_edge)
                if norm:
                    out.append((f"figure {len(out) + 1} embedded in {doc_name}", norm[0], norm[1]))
            if len(out) >= max_images:
                break
    elif ext == ".pdf":
        from pypdf import PdfReader
        for pno, page in enumerate(PdfReader(io.BytesIO(data)).pages, 1):
            for img in page.images:
                norm = normalise_image(img.data, "image/" + (img.name.rsplit(".", 1)[-1].lower() or "png"), max_edge)
                if norm:
                    out.append((f"figure {len(out) + 1} embedded in {doc_name} (page {pno})", norm[0], norm[1]))
                if len(out) >= max_images:
                    return out
    return out


def vsdx_dict(data: bytes, name: str, page: str | None = None) -> dict:
    """Parse a .vsdx with the visio-reader skill script (it needs a real path -> temp file).
    `page` restricts the parse to one named page (a `#page` fragment on `name` is honoured too)."""
    from lab.core.visio.read_vsdx import read_vsdx as _rv   # local: the vsdx library loads only when a .vsdx is parsed
    base, frag = split_fragment(name)
    page = page or frag
    # A private directory per call: storage-mcp serves concurrent reads of the SAME workbook (one per
    # page), so a path keyed by pid + file name collided (truncated parses / FileNotFoundError).
    # The file keeps its own name inside the directory (the parser reports `file` from the path).
    with tempfile.TemporaryDirectory(prefix="vsdx-") as d:
        p = Path(d) / (base.rstrip("/").split("/")[-1] or "diagram.vsdx")
        p.write_bytes(data)
        return _rv(str(p), page=page)


def vsdx_page_image(data: bytes, name: str, page: str | None = None,
                    max_edge: int = MAX_IMAGE_EDGE, render=None):
    """A .vsdx PAGE as a picture — the second representation of the same diagram, for a vision model.

    The page NAME (or a `#page` fragment on `name`, the same selector `vsdx_dict` honours) is resolved
    against the workbook's own page order and handed to the renderer as an index; the bitmap comes
    back already through `normalise_image`. Returns None when the page normalises below the sizing
    floor. Raises `ValueError` for an unknown page name and `RuntimeError` when this host has no
    LibreOffice / PDF rasteriser — the caller degrades to the structured parse alone.
    `render` injects the renderer (default `lab.platform.render_vsdx.render_page`)."""
    from lab.core.visio.read_vsdx import page_index, page_names   # local: vsdx loads only for a .vsdx
    base, frag = split_fragment(name)
    page = page or frag
    if render is None:
        from lab.platform import render_vsdx
        render = render_vsdx.render_page
    with tempfile.TemporaryDirectory(prefix="vsdx-") as d:
        p = Path(d) / (base.rstrip("/").split("/")[-1] or "diagram.vsdx")
        p.write_bytes(data)
        index = page_index(page_names(str(p)), page)
    return render(data, index, max_edge=max_edge)
