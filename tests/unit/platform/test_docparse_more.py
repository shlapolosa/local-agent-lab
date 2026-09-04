"""src/lab/platform/docparse.py — document text, the image sizing contract and embedded-figure extraction,
all OFFLINE on fixtures built in-test (Pillow images, python-docx documents, pypdf/Pillow PDFs)."""
import io
import os

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from PIL import Image  # noqa: E402

from lab.platform import docparse  # noqa: E402

FIXTURE = os.path.join(ROOT, "var", "inputs", "visio_to_archimate", "malaffi-application-solution-arch.vsdx")


def _noise(w, h, fmt="PNG", seed=1):
    """An incompressible image (every pixel differs) so even small ones exceed MIN_IMAGE_BYTES."""
    import random
    rnd = random.Random(seed)
    im = Image.new("RGB", (w, h))
    im.putdata([(rnd.randrange(256), rnd.randrange(256), rnd.randrange(256)) for _ in range(w * h)])
    buf = io.BytesIO(); im.save(buf, fmt); return buf.getvalue()


def _flat(w, h, fmt="PNG"):
    buf = io.BytesIO(); Image.new("RGB", (w, h), (10, 20, 30)).save(buf, fmt); return buf.getvalue()


def _docx(paragraphs=(), table=None, pictures=()):
    import docx
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table:
        t = d.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
    for pic in pictures:
        d.add_picture(io.BytesIO(pic))
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


def _text_pdf(text):
    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, StreamObject
    w = PdfWriter(); p = w.add_blank_page(300, 300)
    font = DictionaryObject({NameObject("/Type"): NameObject("/Font"), NameObject("/Subtype"): NameObject("/Type1"),
                             NameObject("/BaseFont"): NameObject("/Helvetica")})
    p[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): w._add_object(font)})})
    s = StreamObject(); s._data = f"BT /F1 12 Tf 20 100 Td ({text}) Tj ET".encode()
    p[NameObject("/Contents")] = w._add_object(s)
    w.add_blank_page(300, 300)                                  # a page with no text -> "" branch
    out = io.BytesIO(); w.write(out); return out.getvalue()


def _image_pdf(*images):
    """Pillow writes each image as a page XObject — what pypdf's page.images reads back."""
    ims = [Image.open(io.BytesIO(b)).convert("RGB") for b in images]
    buf = io.BytesIO(); ims[0].save(buf, "PDF", save_all=True, append_images=ims[1:]); return buf.getvalue()


# ------------------------------------------------------------------ document_text
def test_document_text_docx_pdf_and_plain():
    data = _docx(["Scope", "  ", "The system SHALL log in users."], table=[["Actor", "Role"], ["Clerk", "Submits"]])
    text = docparse.document_text(data, ".docx")
    assert text.splitlines() == ["Scope", "The system SHALL log in users.", "Actor | Role", "Clerk | Submits"]
    pdf = docparse.document_text(_text_pdf("Hello requirements"), ".pdf")
    assert pdf.startswith("Hello requirements") and pdf.endswith("\n\n"), repr(pdf)
    assert docparse.document_text("# Title\nbody".encode(), ".md") == "# Title\nbody"
    assert docparse.document_text(b"caf\xc3\xa9 \xff", ".txt") == "café �", "bad bytes never raise"


def test_document_text_truncates():
    out = docparse.document_text(b"x" * 100, ".txt", max_chars=10)
    assert out.startswith("x" * 10 + "\n\n[truncated: 100 chars in total, first 10 shown]")
    assert docparse.document_text(b"x" * 10, ".txt", max_chars=10) == "x" * 10


# ------------------------------------------------------------------ normalise_image: the sizing contract
def test_normalise_image_rejects_decorations():
    assert docparse.normalise_image(_flat(400, 400), "image/png") is None, "< MIN_IMAGE_BYTES: a bullet/icon"
    small = _noise(40, 40)
    assert len(small) >= docparse.MIN_IMAGE_BYTES and docparse.normalise_image(small, "image/png") is None, "edge < 64 px"
    assert docparse.normalise_image(b"\x00" * 4096, "image/png") is None, "unreadable bytes -> skipped, not raised"


def test_normalise_image_passthrough_and_conversion():
    png = _noise(200, 120)
    assert docparse.normalise_image(png, "image/png") == (png, "image/png", (200, 120)), "in-bounds PNG untouched"
    jpg = _noise(200, 120, "JPEG")
    assert docparse.normalise_image(jpg, "image/jpeg") == (jpg, "image/jpeg", (200, 120))
    gif = _noise(200, 120, "GIF")
    out, mt, size = docparse.normalise_image(gif, "image/gif")
    assert mt == "image/png" and size == (200, 120) and out[:8] == b"\x89PNG\r\n\x1a\n", "other formats -> PNG"
    big = _noise(2000, 500)
    out, mt, size = docparse.normalise_image(big, "image/png")
    assert mt == "image/png" and size == (1600, 400), "downscaled to the 1600 px edge, aspect kept"
    assert Image.open(io.BytesIO(out)).size == (1600, 400)
    out, mt, size = docparse.normalise_image(_noise(300, 100), "image/png", max_edge=150)
    assert size == (150, 50)


# ------------------------------------------------------------------ embedded_figures
def test_embedded_figures_docx():
    figs = docparse.embedded_figures(_docx(["t"], pictures=[_noise(200, 150), _flat(20, 20), _noise(120, 90, "JPEG")]),
                                     ".docx", "req.docx")
    assert [f[0] for f in figs] == ["figure 1 embedded in req.docx", "figure 2 embedded in req.docx"], "decoration dropped"
    assert {f[2] for f in figs} == {"image/png", "image/jpeg"}
    assert Image.open(io.BytesIO(figs[0][1])).size == (200, 150)
    assert len(docparse.embedded_figures(_docx(pictures=[_noise(100, 100), _noise(100, 100)]), ".docx", "r.docx",
                                         max_images=1)) == 1
    assert docparse.embedded_figures(_docx(["no pictures"]), ".docx", "r.docx") == []
    assert docparse.embedded_figures(b"irrelevant", ".md", "r.md") == []


def test_embedded_figures_pdf():
    pdf = _image_pdf(_noise(200, 150), _flat(20, 20), _noise(160, 100))
    figs = docparse.embedded_figures(pdf, ".pdf", "req.pdf")
    assert [f[0] for f in figs] == ["figure 1 embedded in req.pdf (page 1)", "figure 2 embedded in req.pdf (page 3)"], \
        "the 20 px decoration on page 2 is dropped"
    assert all(f[2] in ("image/png", "image/jpeg") for f in figs)
    assert len(docparse.embedded_figures(pdf, ".pdf", "req.pdf", max_images=1)) == 1
    assert docparse.embedded_figures(_text_pdf("no images"), ".pdf", "t.pdf") == []


# ------------------------------------------------------------------ vsdx_dict with a page fragment
def test_vsdx_dict_page_fragment():
    if not os.path.exists(FIXTURE):
        print("  (skip: fixture .vsdx not present)"); return
    data = open(FIXTURE, "rb").read()
    whole = docparse.vsdx_dict(data, "m.vsdx")
    pages = sorted({s.get("page") for s in whole["shapes"] if s.get("page")})
    if not pages:
        print("  (skip: parser reports no page names)"); return
    one = docparse.vsdx_dict(data, f"art://x/m.vsdx#{pages[0]}")
    assert one["file"] == "m.vsdx" and 0 < len(one["shapes"]) <= len(whole["shapes"])
    assert {s.get("page") for s in one["shapes"]} <= {pages[0]}
    assert docparse.vsdx_dict(data, "m.vsdx", page=pages[0])["shapes"] == one["shapes"], "explicit page == fragment"


if __name__ == "__main__":
    for name, fn in list(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn(); print("ok", name)
