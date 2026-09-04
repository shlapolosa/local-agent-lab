"""src/lab/substrate/mcp/storage/server.py — the read-only upload-store tools through an in-memory fastmcp
Client, OFFLINE: a temp LocalStore holds a generated .vsdx (authored with the lab's own
make_sample_vsdx building blocks), a Pillow PNG, a python-docx .docx with an embedded figure and
a table, and a .md. Asserts the fastmcp gotcha (image tools carry NO outputSchema and return image
content blocks), the ref/kind errors, that the tools read the container's UPLOAD store (not the
artifact store) and the `__main__` bucket env check.
Run: .venv/bin/python tests/unit/substrate/mcp/storage/test_server.py   (also pytest-compatible)"""
import asyncio
import importlib.util
import io
import os
import random
import runpy
import sys
import tempfile
import zipfile
from contextlib import contextmanager

import docx
import pytest
from docx.shared import Inches
from fastmcp import Client
from mcp.types import ImageContent, TextContent
from PIL import Image

from lab.platform import config
from lab.substrate import artifacts
from lab.workloads.visio_to_archimate import make_sample_vsdx as MV

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))))
SERVER = os.path.join(ROOT, "src", "lab", "substrate", "mcp", "storage", "server.py")

TMP = srv = UP = REFS = None        # set up by `_server` (never at import: it pins the environment)
ENV_OWNED = False                   # true when no UPLOADS_URL was exported -> we own the fallback

TOOLS = {"storage_list", "storage_info", "storage_get", "storage_read_document", "storage_read_vsdx",
         "storage_extract_figures", "storage_render_vsdx"}
IMAGE_TOOLS = {"storage_get", "storage_extract_figures", "storage_render_vsdx"}


# ---------------------------------------------------------------- fixtures (all generated)
def _png(size=256, fmt="PNG") -> bytes:
    """A noisy image: > MIN_IMAGE_BYTES and >= MIN_IMAGE_EDGE so it counts as a figure."""
    rnd = random.Random(size)
    im = Image.frombytes("RGB", (size, size), bytes(rnd.getrandbits(8) for _ in range(size * size * 3)))
    buf = io.BytesIO(); im.save(buf, fmt); return buf.getvalue()


def _vsdx() -> bytes:
    page_xml, _ = MV.build_page_xml(
        [{"id": "a", "type": "ApplicationComponent", "name": "Portal"},
         {"id": "b", "type": "DataObject", "name": "Claims DB"},
         {"id": "c", "type": "BusinessActor", "name": "Clerk"}],
        [{"type": "Access", "src": "a", "tgt": "b"}, {"type": "Serving", "src": "a", "tgt": "c"}])
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in MV._parts(page_xml).items():
            z.writestr(name, data)
    return buf.getvalue()


def _docx(with_figure=True) -> bytes:
    d = docx.Document()
    d.add_paragraph("Requirements for the claims portal.")
    d.add_paragraph("")                                              # blank paragraphs are dropped
    t = d.add_table(rows=1, cols=2); t.rows[0].cells[0].text = "REQ-1"; t.rows[0].cells[1].text = "Sub-second lookup"
    if with_figure:
        d.add_picture(io.BytesIO(_png()), width=Inches(2))
        d.add_picture(io.BytesIO(_png(8)), width=Inches(0.1))        # a decoration: too small, dropped
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


@pytest.fixture(scope="module", autouse=True)
def _server():
    """Compose the server against a temp LocalStore standing in for the upload store. The server
    composes at import and `lab.platform.config` reads the env once at ITS import, so the pins live
    HERE (a module fixture) rather than at this module's import, where they would leak — most
    sharply `MCP_SHARED_SECRET` and a popped `DATABASE_URL`. Undone with the module."""
    global TMP, srv, UP, REFS, ENV_OWNED
    mp = pytest.MonkeyPatch()
    TMP = tempfile.mkdtemp(prefix="storage-mcp-test-")
    ENV_OWNED = "UPLOADS_URL" not in os.environ           # we control the fallback config resolved
    mp.setenv("MCP_SHARED_SECRET", "shh")                 # the upload store is overridden on the container
    for k in ("OTEL_EXPORTER_OTLP_ENDPOINT", "UPLOADS_URL", "DATABASE_URL",
              "S3_ENDPOINT", "S3_ACCESS_KEY_ID", "S3_SECRET_ACCESS_KEY"):
        mp.delenv(k, raising=False)

    spec = importlib.util.spec_from_file_location("storage_mcp_server", SERVER)
    srv = importlib.util.module_from_spec(spec)
    sys.modules["storage_mcp_server"] = srv
    spec.loader.exec_module(srv)

    UP = artifacts.LocalStore(os.path.join(TMP, "uploads"))
    srv.server.container.uploads.override(UP)        # the UPLOAD store the tools read, via the kit
    REFS = {
        "vsdx": UP.put("lab-system.vsdx", _vsdx(), artifacts.content_type_for("x.vsdx")),
        "png": UP.put("diagram.png", _png(), "image/png"),
        "jpg": UP.put("photo.jpg", _png(300, "JPEG"), "image/jpeg"),
        "big": UP.put("big.png", _png(700), "image/png"),
        "tiny": UP.put("icon.png", _png(16), "image/png"),
        "docx": UP.put("req.docx", _docx(), artifacts.content_type_for("x.docx")),
        "plain": UP.put("plain.docx", _docx(with_figure=False), artifacts.content_type_for("x.docx")),
        "md": UP.put("notes.md", b"# Notes\n\nThe portal must be fast.\n", "text/markdown"),
    }
    yield
    sys.modules.pop("storage_mcp_server", None)
    mp.undo()
    TMP = srv = UP = REFS = None


def call(_tool, **args):
    async def go():
        async with Client(srv.server.mcp) as c:
            return await c.call_tool(_tool, args)
    return asyncio.run(go())


def call_error(_tool, **args) -> str:
    async def go():
        async with Client(srv.server.mcp) as c:
            r = await c.call_tool(_tool, args, raise_on_error=False)
            assert r.is_error, f"{_tool} should have failed"
            return r.content[0].text
    return asyncio.run(go())


def tools():
    async def go():
        async with Client(srv.server.mcp) as c:
            return await c.list_tools()
    return asyncio.run(go())


@contextmanager
def uploads_url(url):
    old = srv.config.UPLOADS_URL
    srv.config.UPLOADS_URL = url
    try:
        yield
    finally:
        srv.config.UPLOADS_URL = old


# ---------------------------------------------------------------- tests
def test_catalogue_and_image_tools_have_no_output_schema():
    by = {t.name: t for t in tools()}
    assert set(by) == TOOLS
    for name in IMAGE_TOOLS:                                       # the fastmcp gotcha
        assert by[name].outputSchema is None, name
    for name in TOOLS - IMAGE_TOOLS:
        assert by[name].outputSchema, name


def test_list_and_info():
    items = call("storage_list").data
    assert {i["name"] for i in items} == {"lab-system.vsdx", "diagram.png", "photo.jpg", "big.png", "icon.png",
                                          "req.docx", "plain.docx", "notes.md"}
    assert all(i["ref"].startswith("art://") and i["size"] > 0 for i in items)
    assert {i["name"] for i in call("storage_list", prefix="req").data} == {"req.docx"}
    assert len(call("storage_list", limit=2).data) == 2
    info = call("storage_info", ref=REFS["vsdx"]).data
    assert info["name"] == "lab-system.vsdx" and info["kind"] == "vsdx"
    assert info["content_type"] == "application/vnd.ms-visio.drawing.main+xml" and info["size"] > 0
    assert call("storage_info", ref=REFS["png"]).data["kind"] == "image"
    assert call("storage_info", ref=REFS["md"]).data["kind"] == "document"
    assert "not an artifact ref" in call_error("storage_info", ref="/tmp/x.png")
    assert call_error("storage_info", ref="art://nope/x.png")


def test_get_image_normalised_and_errors():
    r = call("storage_get", ref=REFS["png"])
    assert r.structured_content is None
    img, label = r.content
    assert isinstance(img, ImageContent) and img.mimeType == "image/png" and len(img.data) > 100
    assert isinstance(label, TextContent) and label.text == "diagram.png 256x256 image/png"
    jpg = call("storage_get", ref=REFS["jpg"]).content
    assert jpg[0].mimeType == "image/jpeg" and jpg[1].text == "photo.jpg 300x300 image/jpeg"
    small = call("storage_get", ref=REFS["big"], max_edge=100).content        # downscaled to the contract
    assert small[1].text == "big.png 100x100 image/png"
    assert "is not an image" in call_error("storage_get", ref=REFS["docx"])
    assert "too small" in call_error("storage_get", ref=REFS["tiny"])
    assert call_error("storage_get", ref="art://nope/x.png")


def test_read_document_docx_md_and_errors():
    text = call("storage_read_document", ref=REFS["docx"]).data
    assert "Requirements for the claims portal." in text and "REQ-1 | Sub-second lookup" in text
    assert "\n\n" not in text.split("REQ-1")[0]                                  # blank paragraph dropped
    md = call("storage_read_document", ref=REFS["md"]).data
    assert md.startswith("# Notes")
    cut = call("storage_read_document", ref=REFS["md"], max_chars=10).data
    assert cut.startswith("# Notes") and "[truncated:" in cut and "first 10 shown" in cut
    assert "is not a document" in call_error("storage_read_document", ref=REFS["png"])
    assert "is not a document" in call_error("storage_read_document", ref=REFS["vsdx"])


def test_read_vsdx_pages_and_fragment():
    d = call("storage_read_vsdx", ref=REFS["vsdx"]).data
    assert d["pages"] == ["Lab System"] and d["file"] == "lab-system.vsdx"
    assert {s["text"] for s in d["shapes"]} == {"Portal", "Claims DB", "Clerk"}
    assert {s["master"] for s in d["shapes"]} == {"Component", "Data Store", "Actor"}
    assert len(d["connectors"]) == 2
    labels = {(c["label"], c["from"], c["to"]) for c in d["connectors"]}
    assert labels == {("Access", "Portal", "Claims DB"), ("Serving", "Portal", "Clerk")}
    assert all(c["from_id"] and c["to_id"] and c["page"] == "Lab System" for c in d["connectors"])
    one = call("storage_read_vsdx", ref=REFS["vsdx"], page="Lab System").data
    assert one["page"] == "Lab System" and len(one["shapes"]) == 3
    frag = call("storage_read_vsdx", ref=REFS["vsdx"] + "#Lab System").data
    assert frag["page"] == "Lab System" and len(frag["shapes"]) == 3
    other = call("storage_read_vsdx", ref=REFS["vsdx"], page="Nope").data
    assert other["pages"] == ["Lab System"] and other["shapes"] == []            # enumerated, not parsed
    assert "is not a .vsdx" in call_error("storage_read_vsdx", ref=REFS["png"])
    assert "is not a .vsdx" in call_error("storage_read_vsdx", ref=REFS["docx"] + "#p1")


# ---------------------------------------------------------------- vsdx page -> image
@contextmanager
def fake_renderer(fn):
    """Stand in for LibreOffice + a PDF rasteriser: the whole render hop, injected — including the
    capability gate, since a host with a renderer is by definition capable."""
    old_fn, old_av = srv.docparse.vsdx_page_image, srv.render_vsdx.available
    srv.docparse.vsdx_page_image, srv.render_vsdx.available = fn, lambda: True
    try:
        yield
    finally:
        srv.docparse.vsdx_page_image, srv.render_vsdx.available = old_fn, old_av


def test_render_vsdx_returns_an_image_block_and_a_label():
    seen = {}

    def render(data, name, page=None, **kw):
        seen["name"], seen["page"], seen["bytes"] = name, page, len(data)
        return _png(300), "image/png", (300, 300)

    with fake_renderer(render):
        r = call("storage_render_vsdx", ref=REFS["vsdx"], page="Lab System")
    assert isinstance(r.content[0], ImageContent) and isinstance(r.content[1], TextContent)
    assert "lab-system.vsdx" in r.content[1].text and "300x300" in r.content[1].text
    assert "Lab System" in r.content[1].text
    assert seen["name"] == "lab-system.vsdx" and seen["page"] == "Lab System" and seen["bytes"] > 0


def test_render_vsdx_honours_a_page_fragment_and_defaults_to_the_first_page():
    def render(data, name, page=None, **kw):
        return _png(300), "image/png", (300, 300)

    with fake_renderer(render):
        assert "page 1" in call("storage_render_vsdx", ref=REFS["vsdx"]).content[1].text
        frag = call("storage_render_vsdx", ref=REFS["vsdx"] + "#Lab System")
    assert "Lab System" in frag.content[1].text


def test_render_vsdx_errors_are_explicit_about_kind_blankness_and_capability():
    assert "is not a .vsdx" in call_error("storage_render_vsdx", ref=REFS["png"])
    with fake_renderer(lambda *a, **k: None):
        assert "nothing renderable" in call_error("storage_render_vsdx", ref=REFS["vsdx"])

    def missing(*a, **k):
        raise RuntimeError("LibreOffice (soffice) not found")

    with fake_renderer(missing):
        assert "LibreOffice" in call_error("storage_render_vsdx", ref=REFS["vsdx"])


def test_render_vsdx_names_a_missing_HOST_CAPABILITY_before_it_reads_anything():
    """A deployment without LibreOffice must say exactly that — a caller has to be able to tell a
    missing capability from a missing grant or a broken file, and degrade knowingly."""
    old, srv.render_vsdx.available = srv.render_vsdx.available, lambda: False
    try:
        err = call_error("storage_render_vsdx", ref=REFS["vsdx"])
    finally:
        srv.render_vsdx.available = old
    assert "cannot render" in err and "LibreOffice" in err


def test_read_vsdx_reports_the_geometric_recovery_it_had_to_do():
    """A Lucidchart-shaped parse carries `recovery`; a native one does not, and the span records the
    lines that yielded no link so a partial recovery is auditable, not invisible."""
    d = call("storage_read_vsdx", ref=REFS["vsdx"]).data
    assert "recovery" not in d              # the generated fixture is native Visio


def test_extract_figures():
    r = call("storage_extract_figures", ref=REFS["docx"])
    assert r.structured_content is None
    assert len(r.content) == 2                                                   # the decoration was dropped
    img, label = r.content
    assert isinstance(img, ImageContent) and img.mimeType == "image/png"
    assert label.text == "figure 1 embedded in req.docx"
    assert call("storage_extract_figures", ref=REFS["docx"], max_edge=50).content[0].mimeType == "image/png"
    assert call("storage_extract_figures", ref=REFS["plain"]).content == []
    assert call("storage_extract_figures", ref=REFS["md"]).content == []
    assert "is not a document" in call_error("storage_extract_figures", ref=REFS["vsdx"])


def test_tools_read_the_containers_upload_store():
    if ENV_OWNED:
        assert config.UPLOADS_URL == config.ARTIFACTS_URL                        # UPLOADS_URL unset -> same store
    c = srv.server.container
    assert c.config.uploads_url() == config.UPLOADS_URL and srv.server.uploads is c.uploads
    assert srv.server.uploads() is UP, "the override is what the tools see"
    bucket = artifacts.LocalStore(f"{TMP}/bucket")                               # a different, empty store
    with c.uploads.override(bucket):
        assert srv.server.uploads() is bucket and call("storage_list").data == []
    assert srv.server.uploads() is UP
    assert srv._name("art://ab12/dir/x.png/") == "x.png"


def test_main_bucket_env_check_and_serve():
    import lab.substrate.mcpserver as ms
    served = []
    real = ms.serve
    ms.serve = lambda mcp, service, port, **kw: served.append((service, port))
    try:
        runpy.run_path(SERVER, run_name="__main__")                              # file:// store: no S3 needed
        assert served == [("storage-mcp", config.STORAGE_MCP_PORT)]
        with uploads_url("s3://bucket/prefix"):
            try:
                runpy.run_path(SERVER, run_name="__main__")
            except SystemExit as e:
                assert "S3_ENDPOINT" in str(e) and "bucket" in str(e)
            else:
                raise AssertionError("a bucket without S3_* credentials must exit")
            os.environ.update({"S3_ENDPOINT": "https://s3.test", "S3_ACCESS_KEY_ID": "k", "S3_SECRET_ACCESS_KEY": "s"})
            try:
                runpy.run_path(SERVER, run_name="__main__")
            finally:
                for k in ("S3_ENDPOINT", "S3_ACCESS_KEY_ID", "S3_SECRET_ACCESS_KEY"):
                    os.environ.pop(k)
        with uploads_url("s3://user:pw@bucket"):                                  # credentials masked in the banner
            os.environ.update({"S3_ENDPOINT": "e", "S3_ACCESS_KEY_ID": "k", "S3_SECRET_ACCESS_KEY": "s"})
            try:
                runpy.run_path(SERVER, run_name="__main__")
            finally:
                for k in ("S3_ENDPOINT", "S3_ACCESS_KEY_ID", "S3_SECRET_ACCESS_KEY"):
                    os.environ.pop(k)
        assert len(served) == 3
    finally:
        ms.serve = real


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-q"]))
