"""Deterministic parsing of a workload's INPUTS — bytes in, plain data out. No I/O, no egress.

Shared by two callers with different trust: `mcp/storage_mcp/server.py` (the governed path — an
agent asks the gateway, the gateway asks storage-mcp, storage-mcp reads the bucket and parses here)
and `processes/visio_to_archimate/inputs.py` (local-dev paths only). Keeping the parsing here means
both produce byte-identical results and the sizing contract the visio-reader skill documents is
enforced in exactly one place.
"""
from __future__ import annotations

import io
import os
import sys
import tempfile
from pathlib import Path

IMAGE_TYPES = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
               ".gif": "image/gif", ".webp": "image/webp"}
DOC_TYPES = {".docx", ".pdf", ".md", ".markdown", ".txt", ".rst", ".csv"}
MAX_DOC_CHARS = int(os.environ.get("BA_MAX_DOC_CHARS", "60000"))      # keeps one doc inside a turn
MAX_EMBEDDED_IMAGES = int(os.environ.get("BA_MAX_EMBEDDED_IMAGES", "8"))
MIN_IMAGE_BYTES = 2048            # below this it is a bullet/icon, not a figure
MIN_IMAGE_EDGE = 64
MAX_IMAGE_EDGE = 1600             # the sizing contract: no image reaches a model larger than this


def kind(name: str) -> str:
    """'vsdx' | 'image' | 'document' | 'unknown' — decided by the file name (ref or path alike)."""
    ext = Path(name.rstrip("/").split("/")[-1]).suffix.lower()
    if ext == ".vsdx":
        return "vsdx"
    if ext in IMAGE_TYPES:
        return "image"
    if ext in DOC_TYPES:
        return "document"
    return "unknown"


def media_type(name: str) -> str:
    return IMAGE_TYPES.get(Path(name.rstrip("/").split("/")[-1]).suffix.lower(), "application/octet-stream")


def ext_of(name: str) -> str:
    return Path(name.rstrip("/").split("/")[-1]).suffix.lower()


def document_text(data: bytes, ext: str, max_chars: int = MAX_DOC_CHARS) -> str:
    """Requirements document -> plain text (.docx paragraphs + tables, .pdf page text, else UTF-8)."""
    if ext == ".docx":
        import docx
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        text = "\n".join(parts)
    elif ext == ".pdf":
        from pypdf import PdfReader
        text = "\n\n".join((pg.extract_text() or "") for pg in PdfReader(io.BytesIO(data)).pages)
    else:
        text = data.decode("utf-8", errors="replace")
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[truncated: {len(text)} chars in total, first {max_chars} shown]"
    return text


def normalise_image(data: bytes, ctype: str, max_edge: int = MAX_IMAGE_EDGE):
    """Apply the sizing contract: reject decorations (< MIN_IMAGE_BYTES or < MIN_IMAGE_EDGE px),
    downscale to max_edge, emit PNG unless already PNG/JPEG within bounds.
    Returns (bytes, media_type, (w, h)) or None when the input is not a usable figure."""
    if len(data) < MIN_IMAGE_BYTES:
        return None
    try:
        from PIL import Image
        im = Image.open(io.BytesIO(data))
        if min(im.size) < MIN_IMAGE_EDGE:
            return None
        if max(im.size) > max_edge or ctype not in ("image/png", "image/jpeg"):
            im.thumbnail((max_edge, max_edge))
            buf = io.BytesIO()
            im.convert("RGB").save(buf, "PNG")
            return buf.getvalue(), "image/png", im.size
        return data, ctype, im.size
    except Exception:
        return None                  # emf/wmf and friends — PIL cannot read them; skip


def embedded_figures(data: bytes, ext: str, doc_name: str, max_images: int = MAX_EMBEDDED_IMAGES,
                     max_edge: int = MAX_IMAGE_EDGE) -> list[tuple[str, bytes, str]]:
    """Figures EMBEDDED in a document -> [(label, bytes, media_type)]: a .docx's image parts, a
    .pdf's page images — normalised so a vision model can read them as diagrams/screenshots."""
    out: list[tuple[str, bytes, str]] = []
    if ext == ".docx":
        import docx
        d = docx.Document(io.BytesIO(data))
        for rel in d.part.rels.values():
            if "image" in rel.reltype and getattr(rel, "target_part", None) is not None:
                norm = normalise_image(rel.target_part.blob, rel.target_part.content_type, max_edge)
                if norm:
                    out.append((f"figure {len(out) + 1} embedded in {doc_name}", norm[0], norm[1]))
            if len(out) >= max_images:
                break
    elif ext == ".pdf":
        from pypdf import PdfReader
        for pno, page in enumerate(PdfReader(io.BytesIO(data)).pages, 1):
            for img in page.images:
                norm = normalise_image(img.data, "image/" + (img.name.rsplit(".", 1)[-1].lower() or "png"), max_edge)
                if norm:
                    out.append((f"figure {len(out) + 1} embedded in {doc_name} (page {pno})", norm[0], norm[1]))
                if len(out) >= max_images:
                    return out
    return out


def vsdx_dict(data: bytes, name: str) -> dict:
    """Parse a .vsdx with the visio-reader skill script (it needs a real path -> temp file)."""
    root = Path(__file__).resolve().parents[1]
    scripts = str(root / ".claude" / "skills" / "visio-reader" / "scripts")
    if scripts not in sys.path:
        sys.path.insert(0, scripts)
    from read_vsdx import read_vsdx as _rv  # noqa: E402
    p = Path(tempfile.gettempdir()) / f"{os.getpid()}-{name.rstrip('/').split('/')[-1]}"
    p.write_bytes(data)
    try:
        return _rv(str(p))
    finally:
        try:
            p.unlink()
        except OSError:
            pass
